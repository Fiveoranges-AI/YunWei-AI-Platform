"""DOCX parser using ``python-docx``.

Walks paragraphs and tables in document order, emitting one chunk per
paragraph and one chunk per table cell. Source ids follow the stable
forms ``docx:p<n>`` and ``docx:table<n>:R<row>C<col>`` so downstream
extraction can ground extracted fields back to a specific paragraph or
cell.
"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any

from yunwei_win.services.schema_ingest.parse_artifact import (
    ParseArtifact,
    ParseCapabilities,
    ParseChunk,
    ParseTable,
    ParseTableCell,
)


class DocxParser:
    async def parse_file(
        self,
        path: Path,
        *,
        filename: str,
        content_type: str | None,
        source_type: str = "docx",
    ) -> ParseArtifact:
        return await asyncio.to_thread(
            self._parse_sync, path, filename, content_type, source_type
        )

    def _parse_sync(
        self,
        path: Path,
        filename: str,
        content_type: str | None,
        source_type: str,
    ) -> ParseArtifact:
        from docx import Document  # local import: keeps optional dep lazy

        doc = Document(str(path))

        chunks: list[ParseChunk] = []
        tables: list[ParseTable] = []
        grounding: dict[str, Any] = {}
        markdown_parts: list[str] = []

        para_idx = 0
        for para in doc.paragraphs:
            para_idx += 1
            text = para.text or ""
            chunk_id = f"docx:p{para_idx}"
            chunks.append(ParseChunk(id=chunk_id, type="paragraph", text=text))
            grounding[chunk_id] = {"paragraph": para_idx}
            if text.strip():
                markdown_parts.append(text)

        for table_idx, table in enumerate(doc.tables, start=1):
            table_id = f"docx:table{table_idx}"
            cells: list[ParseTableCell] = []
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    cell_text = cell.text or ""
                    cell_id = f"{table_id}:R{r_idx}C{c_idx}"
                    cells.append(
                        ParseTableCell(
                            row=r_idx,
                            col=c_idx,
                            text=cell_text,
                            ref_id=cell_id,
                        )
                    )
                    chunks.append(
                        ParseChunk(id=cell_id, type="table_cell", text=cell_text)
                    )
                    grounding[cell_id] = {
                        "table_id": table_id,
                        "row": r_idx,
                        "col": c_idx,
                    }
            tables.append(ParseTable(id=table_id, cells=cells))
            markdown_parts.append(_table_to_markdown(table_id, cells))

        markdown = "\n\n".join(markdown_parts)

        return ParseArtifact(
            version=1,
            provider="docx",
            source_type=source_type,
            markdown=markdown,
            pages=[{"id": "page:1", "page_number": 1}],
            chunks=chunks,
            grounding=grounding,
            tables=tables,
            metadata={
                "filename": filename,
                "paragraph_count": para_idx,
                "table_count": len(tables),
            },
            capabilities=ParseCapabilities(
                pages=True,
                chunks=True,
                docx_paragraphs=True,
                docx_tables=bool(tables),
                table_cells=bool(tables),
            ),
        )


def _table_to_markdown(table_id: str, cells: list[ParseTableCell]) -> str:
    if not cells:
        return f"<!-- empty {table_id} -->"
    grid: dict[int, dict[int, str]] = {}
    max_row = 0
    max_col = 0
    for cell in cells:
        grid.setdefault(cell.row, {})[cell.col] = cell.text
        max_row = max(max_row, cell.row)
        max_col = max(max_col, cell.col)
    lines: list[str] = []
    for r in range(1, max_row + 1):
        row_cells = grid.get(r, {})
        lines.append(
            "| " + " | ".join(row_cells.get(c, "") for c in range(1, max_col + 1)) + " |"
        )
    return "\n".join(lines)
