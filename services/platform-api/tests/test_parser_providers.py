from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import pytest

from yunwei_win.services.schema_ingest.parsers.docx import DocxParser
from yunwei_win.services.schema_ingest.parsers.landingai import LandingAIParser
from yunwei_win.services.schema_ingest.parsers.spreadsheet import SpreadsheetParser
from yunwei_win.services.schema_ingest.parsers.text import TextParser


@pytest.fixture(autouse=True)
def _clean_state():
    yield


@pytest.mark.asyncio
async def test_text_parser_creates_single_chunk_artifact():
    artifact = await TextParser().parse_text(
        "客户：测试有限公司\n金额：30000", filename="note.txt"
    )
    assert artifact.provider == "text"
    assert artifact.source_type == "text"
    assert artifact.markdown.startswith("客户：")
    assert artifact.chunks[0].id == "text:0"
    assert artifact.capabilities.text_spans is True


@pytest.mark.asyncio
async def test_landingai_parser_preserves_chunks_splits_and_grounding(
    monkeypatch, tmp_path
):
    async def fake_parse(path: Path):
        return SimpleNamespace(
            markdown="# Parsed",
            chunks=[{"id": "c1", "text": "客户"}],
            splits=[{"id": "page:1", "page_number": 1}],
            grounding={"c1": {"page": 1, "bbox": [0, 0, 10, 10]}},
            metadata={"page_count": 1},
        )

    monkeypatch.setattr(
        "yunwei_win.services.schema_ingest.parsers.landingai.parse_file_to_markdown",
        fake_parse,
    )
    path = tmp_path / "contract.pdf"
    path.write_bytes(b"%PDF")

    artifact = await LandingAIParser().parse_file(
        path,
        filename="contract.pdf",
        content_type="application/pdf",
        source_type="pdf",
    )

    assert artifact.provider == "landingai"
    assert artifact.markdown == "# Parsed"
    assert artifact.chunks[0].id == "c1"
    assert artifact.grounding["c1"]["bbox"] == [0, 0, 10, 10]
    assert artifact.capabilities.visual_grounding is True


@pytest.mark.asyncio
async def test_spreadsheet_parser_emits_sheet_cell_refs(tmp_path):
    import openpyxl

    path = tmp_path / "quote.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "报价单"
    ws["A1"] = "客户"
    ws["B1"] = "金额"
    ws["A2"] = "测试有限公司"
    ws["B2"] = 30000
    wb.save(path)

    artifact = await SpreadsheetParser().parse_file(
        path,
        filename="quote.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        source_type="spreadsheet",
    )

    assert "报价单" in artifact.markdown
    ids = {chunk.id for chunk in artifact.chunks}
    assert "sheet:报价单!R2C2" in ids
    assert artifact.capabilities.spreadsheet_cells is True


@pytest.mark.asyncio
async def test_docx_parser_emits_paragraph_and_table_refs(tmp_path):
    from docx import Document

    path = tmp_path / "contacts.docx"
    doc = Document()
    doc.add_paragraph("客户：测试有限公司")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "值"
    table.rows[1].cells[0].text = "金额"
    table.rows[1].cells[1].text = "30000"
    doc.save(path)

    artifact = await DocxParser().parse_file(
        path,
        filename="contacts.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        source_type="docx",
    )

    assert artifact.provider == "docx"
    assert artifact.source_type == "docx"
    assert "客户：测试有限公司" in artifact.markdown
    ids = {chunk.id for chunk in artifact.chunks}
    assert "docx:p1" in ids
    assert "docx:table1:R2C2" in ids
    assert artifact.capabilities.docx_paragraphs is True
    assert artifact.capabilities.docx_tables is True


@pytest.mark.asyncio
async def test_spreadsheet_parser_xls_uses_pandas_xlrd(monkeypatch, tmp_path):
    """``.xls`` reads must go through pandas (which uses xlrd as the engine)
    rather than openpyxl, and emit the same cell-ref shape as ``.xlsx``."""

    import pandas as pd

    path = tmp_path / "legacy.xls"
    path.write_bytes(b"fake-xls-bytes")

    captured: dict[str, object] = {}

    def fake_read_excel(file, sheet_name=None, header=None):
        captured["file"] = file
        captured["sheet_name"] = sheet_name
        captured["header"] = header
        return {
            "Sheet1": pd.DataFrame(
                [["客户", "金额"], ["测试有限公司", 30000]]
            )
        }

    monkeypatch.setattr(pd, "read_excel", fake_read_excel)

    artifact = await SpreadsheetParser().parse_file(
        path,
        filename="legacy.xls",
        content_type="application/vnd.ms-excel",
        source_type="spreadsheet",
    )

    assert str(captured["file"]) == str(path)
    assert captured["sheet_name"] is None  # all sheets
    ids = {chunk.id for chunk in artifact.chunks}
    assert "sheet:Sheet1!R2C2" in ids
    assert artifact.capabilities.spreadsheet_cells is True


def test_xlrd_is_a_declared_dependency():
    """``.xls`` parsing relies on xlrd; guard against the dep silently
    being dropped from pyproject.toml."""

    import importlib

    importlib.import_module("xlrd")
